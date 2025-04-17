import os
import sys
import json
import time
import random
import logging
import argparse
import concurrent.futures
from threading import Lock
from pathlib import Path
from datetime import datetime

import pandas as pd
import requests
from tqdm import tqdm
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class BookGenerator:
    def __init__(self, excel_path, provider=None):
        """Initialize the BookGenerator with the path to the Excel outline.
        
        Args:
            excel_path: Path to the Excel outline file
            provider: API provider to use ('deepseek', 'gemini', 'openrouter', 'siliconflow', 'ark', 'all', or None for random)
        """
        # Load environment variables
        load_dotenv(override=True)  # Force reload and override existing env vars
        
        # Set up logging
        self._setup_logging()
        
        # Available API providers
        self.providers = {
            'deepseek': {
                'api_key': os.getenv('DEEPSEEK_API_KEY'),
                'api_url': os.getenv('DEEPSEEK_API_URL'),
                'model': os.getenv('DEEPSEEK_API_MODEL'),
                'headers': lambda key: {"Authorization": f"Bearer {key}"},
                'request_format': self._format_openai_request,
                'parse_response': self._parse_openai_response,
                'enabled': bool(os.getenv('DEEPSEEK_API_KEY') and os.getenv('DEEPSEEK_API_URL') and os.getenv('DEEPSEEK_API_MODEL')),
                'last_used': 0,  # Timestamp of last successful use
                'success_count': 0,
                'failure_count': 0,
                'consecutive_failures': 0
            },
            'gemini': {
                'api_key': os.getenv('GEMINI_API_KEY'),
                'api_url': os.getenv('GEMINI_API_URL'),
                'model': os.getenv('GEMINI_API_MODEL'),
                'headers': lambda key: {"Content-Type": "application/json"},
                'request_format': self._format_gemini_request,
                'parse_response': self._parse_gemini_response,
                'enabled': bool(os.getenv('GEMINI_API_KEY') and os.getenv('GEMINI_API_URL') and os.getenv('GEMINI_API_MODEL')),
                'last_used': 0,
                'success_count': 0,
                'failure_count': 0,
                'consecutive_failures': 0
            },
            'openrouter': {
                'api_key': os.getenv('OPENROUTER_API_KEY'),
                'api_url': os.getenv('OPENROUTER_API_URL'),
                'model': os.getenv('OPENROUTER_MODEL'),
                'headers': lambda key: {
                    "Authorization": f"Bearer {key}",
                    "HTTP-Referer": "https://github.com/user/repo"
                },
                'request_format': self._format_openai_request,
                'parse_response': self._parse_openai_response,
                'enabled': bool(os.getenv('OPENROUTER_API_KEY') and os.getenv('OPENROUTER_API_URL') and os.getenv('OPENROUTER_MODEL')),
                'last_used': 0,
                'success_count': 0,
                'failure_count': 0,
                'consecutive_failures': 0
            },
            'siliconflow': {
                'api_key': os.getenv('SILICONFLOW_API_KEY'),
                'api_url': os.getenv('SILICONFLOW_API_URL'),
                'model': os.getenv('SILICONFLOW_MODEL'),
                'headers': lambda key: {"Authorization": f"Bearer {key}"},
                'request_format': self._format_openai_request,
                'parse_response': self._parse_openai_response,
                'enabled': bool(os.getenv('SILICONFLOW_API_KEY') and os.getenv('SILICONFLOW_API_URL') and os.getenv('SILICONFLOW_MODEL')),
                'last_used': 0,
                'success_count': 0,
                'failure_count': 0,
                'consecutive_failures': 0
            },
            'ark': {
                'api_key': os.getenv('ARK_API_KEY'),
                'api_url': os.getenv('ARK_API_URL'),
                'model': os.getenv('ARK_MODEL'),
                'headers': lambda key: {"Authorization": f"Bearer {key}"},
                'request_format': self._format_openai_request,
                'parse_response': self._parse_openai_response,
                'enabled': bool(os.getenv('ARK_API_KEY') and os.getenv('ARK_API_URL') and os.getenv('ARK_MODEL')),
                'last_used': 0,
                'success_count': 0,
                'failure_count': 0,
                'consecutive_failures': 0
            }
        }
        
        # Get available providers
        self.available_providers = [p for p, data in self.providers.items() if data['enabled']]
        if not self.available_providers:
            self.logger.error("No API providers available. Please check your .env file.")
            raise ValueError("No API providers available. Please check your .env file.")
        
        self.logger.info(f"Available API providers: {', '.join(self.available_providers)}")
        
        # Select provider mode
        if provider == 'all':
            self.use_all_providers = True
            self.active_provider = None
            self.logger.info("Using all available API providers in parallel")
        else:
            self.use_all_providers = False
            # Select specific provider or random
            if provider and provider in self.providers and self.providers[provider]['enabled']:
                self.active_provider = provider
            else:
                self.active_provider = random.choice(self.available_providers)
            
            # Get selected provider configuration
            provider_config = self.providers[self.active_provider]
            self.api_key = provider_config['api_key']
            self.api_url = provider_config['api_url']
            self.model = provider_config['model']
            self.get_headers = provider_config['headers']
            self.format_request = provider_config['request_format']
            self.parse_response = provider_config['parse_response']
            
            # Debug logging for environment variables
            self.logger.info(f"Selected API provider: {self.active_provider}")
            self.logger.info(f"API URL: {self.api_url}")
            self.logger.info(f"Model: {self.model}")
            self.logger.info(f"API Key: {'Set (' + self.api_key[:5] + '...)' if self.api_key else 'Not set'}")
            
            # Clean up API key format
            self.api_key = self.api_key.strip()  # Remove any whitespace
            if self.api_key.startswith('Bearer '):
                self.api_key = self.api_key[7:]  # Remove Bearer prefix if present
        
        # Initialize other properties
        self.excel_path = excel_path
        self.outline = None
        self.generated_chapters = {}
        self.context = []
        self.doc = Document()
        self._setup_document_style()
        self.api_lock = Lock()  # Lock for API calls to prevent rate limiting issues
        
        # Adjust max_workers based on available providers when using all providers
        if self.use_all_providers:
            self.max_workers = min(len(self.available_providers) * 2, 8)  # Scale with providers but cap at 8
        else:
            self.max_workers = 4  # Default for single provider
            
        self.logger.info(f"Using {self.max_workers} worker threads")
        
        # Optional configuration
        self.temperature = float(os.getenv('TEMPERATURE', '0.7'))
        self.max_tokens = int(os.getenv('MAX_TOKENS', '10000'))
        
    def _setup_logging(self):
        """Set up logging configuration."""
        # Create logs directory if it doesn't exist
        log_dir = Path("logs")
        log_dir.mkdir(exist_ok=True)
        
        # Configure logger
        self.logger = logging.getLogger('BookGenerator')
        self.logger.setLevel(logging.INFO)
        
        # Clear existing handlers if any
        if self.logger.handlers:
            self.logger.handlers.clear()
        
        # Create file handler for logging to file
        log_file = log_dir / f"book_generator_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
        file_handler = logging.FileHandler(log_file, encoding='utf-8')
        file_handler.setLevel(logging.INFO)
        
        # Create console handler for logging to console
        console_handler = logging.StreamHandler()
        console_handler.setLevel(logging.INFO)
        
        # Create formatter and add it to the handlers
        formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
        file_handler.setFormatter(formatter)
        console_handler.setFormatter(formatter)
        
        # Add the handlers to the logger
        self.logger.addHandler(file_handler)
        self.logger.addHandler(console_handler)
        
        self.logger.info("Logging initialized")

    def _format_openai_request(self, prompt, max_tokens):
        """Format request data for OpenAI-compatible APIs (DeepSeek, OpenRouter, Siliconflow, Ark)."""
        return {
            "model": self.model,
            "messages": [
                {"role": "system", "content": "你是一位专业的教材作者，擅长创作清晰、专业的教育内容。"},
                {"role": "user", "content": prompt}
            ],
            "temperature": self.temperature,
            "max_tokens": max_tokens,
            "stream": False
        }
    
    def _format_gemini_request(self, prompt, max_tokens):
        """Format request data for Gemini API."""
        return {
            "contents": [
                {"role": "user", "parts": [{"text": prompt}]}
            ],
            "generationConfig": {
                "temperature": self.temperature,
                "maxOutputTokens": max_tokens
            }
        }
    
    def _parse_gemini_response(self, response_json):
        """Parse response from Gemini API."""
        try:
            return response_json.get('candidates', [{}])[0].get('content', {}).get('parts', [{}])[0].get('text', '')
        except (KeyError, IndexError, TypeError):
            print(f"无法解析Gemini响应: {json.dumps(response_json, ensure_ascii=False)[:200]}...")
            return None
    
    def _parse_openai_response(self, response_json):
        """Parse response from OpenAI-compatible APIs."""
        try:
            return response_json.get('choices', [{}])[0].get('message', {}).get('content', '')
        except (KeyError, IndexError, TypeError):
            print(f"无法解析OpenAI兼容响应: {json.dumps(response_json, ensure_ascii=False)[:200]}...")
            return None

    def _setup_document_style(self):
        """Set up the Word document style."""
        # Set up styles for different heading levels
        style = self.doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
        
        # Title style
        style = self.doc.styles['Title']
        style.font.name = '黑体'
        style.font.size = Pt(24)
        style.font.bold = True
        
        # Heading 1 style
        style = self.doc.styles['Heading 1']
        style.font.name = '黑体'
        style.font.size = Pt(18)
        style.font.bold = True
        
        # Heading 2 style
        style = self.doc.styles['Heading 2']
        style.font.name = '黑体'
        style.font.size = Pt(16)
        style.font.bold = True

        # Add title page
        title = self.doc.add_paragraph("财富管理教材", style='Title')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph()  # Add some space
        self.doc.add_page_break()

    def load_outline(self):
        """Load and validate the book outline from Excel file."""
        try:
            self.outline = pd.read_excel(self.excel_path)
            # Check if the file has the expected Chinese column names
            required_columns = ['课程模块', '课程主题', '课程大纲', '内容说明']
            if not all(col in self.outline.columns for col in required_columns):
                raise ValueError(f"Excel file must contain columns: {required_columns}")
            # Clean any NaN values
            self.outline = self.outline.fillna('')
            return True
        except Exception as e:
            print(f"Error loading outline: {str(e)}")
            return False

    def _make_api_call(self, prompt, max_tokens=8000):
        """Make an API call with consistent headers and error handling."""
        # If using all providers, try them in parallel
        if self.use_all_providers:
            return self._make_parallel_api_calls(prompt, max_tokens)
        
        # Otherwise use the selected provider with fallback
        return self._make_single_api_call(self.active_provider, prompt, max_tokens)
    
    def _make_single_api_call(self, provider_name, prompt, max_tokens, is_fallback=False):
        """Make an API call to a specific provider with error handling."""
        provider_config = self.providers[provider_name]
        api_key = provider_config['api_key']
        api_url = provider_config['api_url']
        model = provider_config['model']
        get_headers = provider_config['headers']
        format_request = provider_config['request_format']
        parse_response = provider_config['parse_response']
        
        headers = get_headers(api_key)
        data = format_request(prompt, max_tokens)
        
        # For Gemini, we need to append the model and generateContent to the URL
        url = api_url
        params = {}
        if provider_name == 'gemini':
            url = f"{api_url}/{model}:generateContent"
            params = {"key": api_key}
            # Remove Authorization header for Gemini
            if "Authorization" in headers:
                del headers["Authorization"]
        
        self.logger.info(f"Making API call to provider: {provider_name}")
        self.logger.debug(f"Request URL: {url}")
        
        try:
            start_time = time.time()
            response = requests.post(url, headers=headers, json=data, params=params, timeout=60)
            elapsed_time = time.time() - start_time
            
            if response.status_code != 200:
                self.logger.error(f"API error ({provider_name}, status {response.status_code}): {response.text[:200]}")
                if response.status_code == 401:
                    self.logger.error(f"Authentication error for {provider_name}. Check API key format.")
                
                # Update provider stats
                provider_config['failure_count'] += 1
                provider_config['consecutive_failures'] += 1
                
                # If not already a fallback call, try another provider
                if not is_fallback:
                    return self._try_fallback_provider(prompt, max_tokens)
                return None
            
            result = response.json()
            content = parse_response(result)
            
            if not content:
                self.logger.warning(f"Empty content returned from {provider_name}")
                
                # Update provider stats
                provider_config['failure_count'] += 1
                provider_config['consecutive_failures'] += 1
                
                # If not already a fallback call, try another provider
                if not is_fallback:
                    return self._try_fallback_provider(prompt, max_tokens)
                return None
            
            # Update provider stats on success
            provider_config['last_used'] = time.time()
            provider_config['success_count'] += 1
            provider_config['consecutive_failures'] = 0
            
            self.logger.info(f"Successful API call to {provider_name} (took {elapsed_time:.2f}s)")
            return content
            
        except Exception as e:
            self.logger.error(f"API call exception ({provider_name}): {str(e)}")
            
            # Update provider stats
            provider_config['failure_count'] += 1
            provider_config['consecutive_failures'] += 1
            
            # If not already a fallback call, try another provider
            if not is_fallback:
                return self._try_fallback_provider(prompt, max_tokens)
            return None
    
    def _select_best_provider(self, exclude_provider=None):
        """Select the best provider based on performance metrics."""
        available_providers = [p for p in self.available_providers if p != exclude_provider]
        if not available_providers:
            return None
        
        # Sort providers by success rate and consecutive failures
        sorted_providers = sorted(
            available_providers,
            key=lambda p: (
                -self.providers[p]['consecutive_failures'],  # Prefer providers with fewer consecutive failures
                -(self.providers[p]['success_count'] / max(1, self.providers[p]['success_count'] + self.providers[p]['failure_count'])),  # Higher success rate
                self.providers[p]['last_used'] or 0  # Prefer providers not used recently
            )
        )
        
        # Log provider selection criteria
        self.logger.info("Provider selection report:")
        for provider in sorted_providers:
            success_rate = self.providers[provider]['success_count'] / max(1, self.providers[provider]['success_count'] + self.providers[provider]['failure_count'])
            self.logger.info(f"  {provider}: 连续失败次数={self.providers[provider]['consecutive_failures']}, "
                           f"成功率={success_rate:.2f}, "
                           f"最近使用时间={self.providers[provider]['last_used']}")
        
        return sorted_providers[0] if sorted_providers else None
        
    def _make_parallel_api_calls(self, prompt, max_tokens):
        """Make parallel API calls to all available providers and return the first successful result."""
        self.logger.info(f"Making parallel API calls to {len(self.available_providers)} providers")
        
        # Use _select_best_provider to get sorted providers
        sorted_providers = self._select_best_provider()
        if not sorted_providers:
            self.logger.error("No providers available for API calls")
            return None
        
        # Use threading to make parallel API calls
        with concurrent.futures.ThreadPoolExecutor(max_workers=len(sorted_providers)) as executor:
            # Submit all API calls
            future_to_provider = {}
            for provider in sorted_providers:
                future = executor.submit(self._make_single_api_call, provider, prompt, max_tokens, True)
                future_to_provider[future] = provider
            
            # Get the first successful result
            for future in concurrent.futures.as_completed(future_to_provider):
                provider = future_to_provider[future]
                try:
                    result = future.result()
                    if result:
                        self.logger.info(f"Using result from {provider}")
                        return result
                except Exception as e:
                    self.logger.error(f"Error in parallel API call to {provider}: {str(e)}")
        
        self.logger.error("All parallel API calls failed")
        return None
    
    def _try_fallback_provider(self, prompt, max_tokens):
        """Try a different provider if the current one fails."""
        # Use _select_best_provider to get the best fallback provider
        fallback_provider = self._select_best_provider(exclude_provider=self.active_provider)
        
        if not fallback_provider:
            self.logger.error("No fallback API providers available")
            return None
        
        self.logger.info(f"Switching to fallback provider: {fallback_provider}")
        
        # Save current provider config to restore later
        original_provider = self.active_provider
        original_api_key = self.api_key
        original_api_url = self.api_url
        original_model = self.model
        original_headers = self.get_headers
        original_format = self.format_request
        original_parse = self.parse_response
        
        # Switch to fallback provider
        provider_config = self.providers[fallback_provider]
        self.active_provider = fallback_provider
        self.api_key = provider_config['api_key']
        self.api_url = provider_config['api_url']
        self.model = provider_config['model']
        self.get_headers = provider_config['headers']
        self.format_request = provider_config['request_format']
        self.parse_response = provider_config['parse_response']
        
        try:
            # Make API call with fallback provider
            return self._make_single_api_call(fallback_provider, prompt, max_tokens, True)
        finally:
            # Restore original provider regardless of outcome
            self.active_provider = original_provider
            self.api_key = original_api_key
            self.api_url = original_api_url
            self.model = original_model
            self.get_headers = original_headers
            self.format_request = original_format
            self.parse_response = original_parse

    def generate_section_content(self, module, topic, outline_point, description):
        """Generate content for a single section using the configured API providers."""
        # Create temp directory if it doesn't exist
        temp_dir = Path("temp/sections")
        temp_dir.mkdir(exist_ok=True, parents=True)
        
        # Generate a safe filename from the topic and outline point
        safe_filename = "".join(x for x in f"{topic}_{outline_point[:30]}" if x.isalnum() or x in (' ', '-', '_')).rstrip()
        temp_file = temp_dir / f"{safe_filename}.md"
        
        # Check if we already have this section generated
        if temp_file.exists():
            self.logger.info(f"Loading section from cache: {outline_point[:30]}...")
            with open(temp_file, 'r', encoding='utf-8') as f:
                content = f.read()
                if content.strip():  # If file is not empty
                    return content

        prompt = f"""请基于以下大纲点写一个详细的小节内容（约800-1000字）：

模块：{module}
主题：{topic}
大纲点：{outline_point}
详细说明：{description}

要求：
1. 内容应该具体且专业
2. 使用清晰的写作风格
3. 包含相关的例子和解释
4. 确保内容完整且自成一体
5. 目标字数800-1000字

请撰写小节内容："""

        try:
            # Use API lock to prevent concurrent API calls
            with self.api_lock:
                self.logger.info(f"Generating section: {outline_point[:30]}...")
                section_content = self._make_api_call(prompt, max_tokens=2000)
                if not section_content:
                    self.logger.error(f"Failed to generate content for section: {outline_point[:30]}...")
                    return None

            # Save the generated content to temp file immediately with markdown formatting
            with open(temp_file, 'w', encoding='utf-8') as f:
                f.write(f"### {outline_point}\n\n")  # Add section heading
                f.write(section_content)
            self.logger.info(f"Section saved to temp file: {temp_file}")
            
            return section_content
        except Exception as e:
            self.logger.error(f"Error generating section content ({outline_point[:30]}...): {str(e)}")
            return None

    def stitch_chapter_content(self, module, topic, sections_content):
        """Stitch together section contents into a coherent chapter."""
        temp_dir = Path("temp/chapters")
        temp_dir.mkdir(exist_ok=True, parents=True)
        
        # Generate a safe filename from the topic
        safe_filename = "".join(x for x in topic if x.isalnum() or x in (' ', '-', '_')).rstrip()
        temp_file = temp_dir / f"{safe_filename}.md"
        
        # Check if we already have this chapter stitched
        if temp_file.exists():
            self.logger.info(f"Loading stitched chapter from cache: {topic}")
            with open(temp_file, 'r', encoding='utf-8') as f:
                content = f.read()
                if content.strip():  # If file is not empty
                    return content

        # Count valid sections
        valid_sections = [content for content in sections_content if content]
        self.logger.info(f"Stitching {len(valid_sections)} sections into chapter: {topic}")
        
        if not valid_sections:
            self.logger.error(f"No valid sections to stitch for chapter: {topic}")
            return None
            
        sections_text = "\n\n".join([f"### 小节内容\n{content}" for content in valid_sections])
        
        prompt = f"""请将以下小节内容整合成一个连贯的章节。

模块：{module}
主题：{topic}

{sections_text}

要求：
1. 为章节添加简短的开篇介绍（200-300字）
2. 在小节之间添加适当的过渡语句
3. 确保整体内容连贯流畅
4. 添加总结性的结束语（200-300字）
5. 保持专业的写作风格

请整合并输出完整章节："""

        try:
            # Use API lock to prevent concurrent API calls
            with self.api_lock:
                self.logger.info(f"Stitching chapter: {topic}")
                start_time = time.time()
                chapter_content = self._make_api_call(prompt, max_tokens=4000)
                elapsed_time = time.time() - start_time
                
                if not chapter_content:
                    self.logger.error(f"Failed to stitch chapter: {topic}")
                    return None
                    
                self.logger.info(f"Successfully stitched chapter: {topic} (took {elapsed_time:.2f}s)")

            # Save the stitched content to temp file immediately with markdown formatting
            with open(temp_file, 'w', encoding='utf-8') as f:
                f.write(f"# {topic}\n\n")  # Add chapter title
                f.write(f"## 模块：{module}\n\n")  # Add module info
                f.write(chapter_content)
            self.logger.info(f"Saved stitched chapter to temp file: {temp_file}")
            
            return chapter_content
        except Exception as e:
            self.logger.error(f"Error stitching chapter content ({topic}): {str(e)}")
            return None

    def generate_sections_parallel(self, module, topic, outline_points, description):
        """Generate multiple sections in parallel using threads."""
        sections_content = []
        
        def generate_section_wrapper(point):
            return self.generate_section_content(module, topic, point, description)
        
        # Create a progress bar
        total_points = len(outline_points)
        completed = 0
        self.logger.info(f"Generating {total_points} sections in parallel for topic: {topic}")
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            # Submit all section generation tasks
            future_to_point = {
                executor.submit(generate_section_wrapper, point): point 
                for point in outline_points
            }
            
            # Process completed sections as they finish
            for future in concurrent.futures.as_completed(future_to_point):
                point = future_to_point[future]
                try:
                    section_content = future.result()
                    completed += 1
                    if section_content:
                        sections_content.append(section_content)
                        self.logger.info(f"Generated section {completed}/{total_points}: {point[:30]}...")
                    else:
                        self.logger.warning(f"Failed to generate section {completed}/{total_points}: {point}")
                except Exception as e:
                    completed += 1
                    self.logger.error(f"Error generating section {completed}/{total_points} ({point}): {str(e)}")
        
        self.logger.info(f"Completed generating {len(sections_content)}/{total_points} sections for topic: {topic}")
        return sections_content

    def generate_chapter_content(self, module, topic, outline, description):
        """Generate content for a chapter by first generating sections in parallel and then stitching them together."""
        # Split the outline into individual points
        outline_points = [point.strip() for point in outline.split('\n') if point.strip()]
        
        self.logger.info(f"Starting chapter generation: {topic} with {len(outline_points)} outline points")
        
        # Generate content for each section in parallel
        sections_content = self.generate_sections_parallel(module, topic, outline_points, description)
        
        if not sections_content:
            self.logger.error(f"Error: All sections failed for chapter {topic}")
            return None
        
        self.logger.info(f"Successfully generated {len(sections_content)}/{len(outline_points)} sections for chapter: {topic}")
        
        # Stitch the sections together into a coherent chapter
        return self.stitch_chapter_content(module, topic, sections_content)

    def generate_chapters_parallel(self, chapters_data):
        """Generate multiple chapters in parallel."""
        generated_content = []
        total_chapters = len(chapters_data)
        completed_chapters = 0
        
        self.logger.info(f"Starting parallel generation of {total_chapters} chapters")
        
        def generate_chapter_wrapper(chapter_info):
            nonlocal completed_chapters
            chapter_number, module, topic, outline, description = chapter_info
            
            try:
                self.logger.info(f"Starting chapter {chapter_number}/{total_chapters}: {topic}")
                start_time = time.time()
                content = self.generate_chapter_content(module, topic, outline, description)
                elapsed_time = time.time() - start_time
                
                completed_chapters += 1
                if content:
                    self.logger.info(f"Completed chapter {chapter_number}/{total_chapters}: {topic} (took {elapsed_time:.2f}s)")
                    return {
                        'chapter_number': chapter_number,
                        'module': module,
                        'topic': topic,
                        'content': content,
                        'success': True
                    }
                else:
                    self.logger.error(f"Failed to generate chapter {chapter_number}/{total_chapters}: {topic}")
                    return {
                        'chapter_number': chapter_number,
                        'module': module,
                        'topic': topic,
                        'content': None,
                        'success': False
                    }
            except Exception as e:
                completed_chapters += 1
                self.logger.error(f"Error generating chapter {chapter_number}: {topic} - {str(e)}")
                return {
                    'chapter_number': chapter_number,
                    'module': module,
                    'topic': topic,
                    'content': None,
                    'success': False
                }
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            # Submit all chapter generation tasks
            future_to_chapter = {
                executor.submit(generate_chapter_wrapper, (
                    idx + 1, row['课程模块'], row['课程主题'], 
                    row['课程大纲'], row['内容说明']
                )): idx 
                for idx, row in enumerate(chapters_data)
            }
            
            # Process completed chapters as they finish
            for future in concurrent.futures.as_completed(future_to_chapter):
                try:
                    chapter_result = future.result()
                    if chapter_result['success']:
                        generated_content.append(chapter_result)
                except Exception as e:
                    self.logger.error(f"Unexpected error processing chapter result: {str(e)}")
        
        # Log summary statistics
        success_count = len(generated_content)
        self.logger.info(f"Chapter generation complete: {success_count}/{total_chapters} chapters successfully generated")
        
        # Sort chapters by chapter number to maintain order
        return sorted(generated_content, key=lambda x: x['chapter_number'])

    def load_outline(self):
        """Load and validate the book outline from Excel file."""
        try:
            self.logger.info(f"Loading outline from Excel file: {self.excel_path}")
            self.outline = pd.read_excel(self.excel_path)
            # Check if the file has the expected Chinese column names
            required_columns = ['课程模块', '课程主题', '课程大纲', '内容说明']
            if not all(col in self.outline.columns for col in required_columns):
                self.logger.error(f"Excel file must contain columns: {required_columns}")
                raise ValueError(f"Excel file must contain columns: {required_columns}")
            # Clean any NaN values
            self.outline = self.outline.fillna('')
            self.logger.info(f"Successfully loaded outline with {len(self.outline)} rows")
            return True
        except Exception as e:
            self.logger.error(f"Error loading outline: {str(e)}")
            return False

    def generate_book(self, output_dir="output"):
        """Generate the entire book with parallel chapter generation."""
        start_time = time.time()
        self.logger.info(f"Starting book generation process")
        
        if not self.load_outline():
            self.logger.error("Failed to load outline, aborting book generation")
            return False

        output_path = Path(output_dir)
        output_path.mkdir(exist_ok=True)
        self.logger.info(f"Output directory: {output_path}")

        # Create table of contents
        self.doc.add_heading('目录', level=1)
        self.doc.add_paragraph()

        # Prepare chapter data for parallel processing
        chapters_data = [
            row for idx, row in self.outline.iterrows()
            if row['课程主题'].strip()  # Skip empty rows
        ]
        
        self.logger.info(f"Found {len(chapters_data)} valid chapters to generate")

        # Generate all chapters in parallel
        self.logger.info("Starting parallel generation of all chapters...")
        generated_content = self.generate_chapters_parallel(chapters_data)

        # Organize and write content
        self.logger.info(f"Writing {len(generated_content)} generated chapters to document")
        for item in generated_content:
            chapter_number = item['chapter_number']
            module = item['module']
            topic = item['topic']
            content = item['content']

            # Add to table of contents
            toc_entry = self.doc.add_paragraph(f"第{chapter_number}章 {topic}")
            toc_entry.style = 'TOC 2'

            # Add chapter content
            self.doc.add_page_break()
            self.doc.add_heading(f'第{chapter_number}章 {topic}', level=1)
            
            if content:
                self.doc.add_paragraph(content)
                
                # Save individual chapter files in markdown format
                chapter_filename = output_path / f"chapter_{chapter_number:02d}.md"
                with open(chapter_filename, 'w', encoding='utf-8') as f:
                    f.write(f"# 第{chapter_number}章: {topic}\n\n")  # Add chapter title
                    f.write(f"## 模块：{module}\n\n")  # Add module info
                    f.write(content)
                
                self.generated_chapters[chapter_number] = {
                    'module': module,
                    'topic': topic,
                    'filename': str(chapter_filename)
                }
                self.logger.info(f"Saved chapter {chapter_number}: {topic}")
            else:
                self.doc.add_paragraph("内容生成失败")
                self.logger.error(f"Failed to generate chapter: {topic}")

        # Save the complete document
        docx_path = output_path / "完整教材.docx"
        self.doc.save(str(docx_path))
        self.logger.info(f"Saved complete document to: {docx_path}")

        # Save metadata
        metadata = {
            'chapters': self.generated_chapters,
            'total_chapters': len(self.generated_chapters),
            'generation_time': time.time() - start_time,
            'timestamp': datetime.now().isoformat(),
            'provider_stats': {name: {
                'success_count': config['success_count'],
                'failure_count': config['failure_count']
            } for name, config in self.providers.items() if config['enabled']}
        }
        metadata_path = output_path / 'book_metadata.json'
        with open(metadata_path, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, indent=2, ensure_ascii=False)
        self.logger.info(f"Saved metadata to: {metadata_path}")
        
        # Log summary statistics
        elapsed_time = time.time() - start_time
        self.logger.info(f"Book generation completed in {elapsed_time:.2f} seconds")
        self.logger.info(f"Successfully generated {len(self.generated_chapters)}/{len(chapters_data)} chapters")
        
        # Log provider statistics
        self.logger.info("Provider statistics:")
        for name, config in self.providers.items():
            if config['enabled']:
                success_rate = config['success_count'] / max(1, config['success_count'] + config['failure_count']) * 100
                self.logger.info(f"  {name}: {config['success_count']} successes, {config['failure_count']} failures ({success_rate:.1f}% success rate)")

        return True

if __name__ == "__main__":
    # Create command line argument parser
    parser = argparse.ArgumentParser(description="生成教育内容书籍")
    parser.add_argument("--excel", "-e", default="data/book_outline.xlsx", 
                        help="Excel文件路径 (默认: data/book_outline.xlsx)")
    parser.add_argument("--provider", "-p", choices=["deepseek", "gemini", "openrouter", "siliconflow", "ark", "all"],
                        help="指定API提供商 (默认: 随机选择一个可用的提供商, 'all'表示使用所有可用提供商)")
    parser.add_argument("--output", "-o", default="output",
                        help="输出目录 (默认: output)")
    parser.add_argument("--workers", "-w", type=int, default=None,
                        help="并行工作线程数 (默认: 根据提供商数量自动设置)")
    
    args = parser.parse_args()
    excel_path = args.excel
    
    # Configure basic logging for startup errors
    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
    logger = logging.getLogger('startup')
    
    # Check if the file exists
    if not os.path.exists(excel_path):
        logger.error(f"File not found: {excel_path}")
        sys.exit(1)
        
    logger.info(f"Using Excel file: {excel_path}")
    logger.info(f"Using API provider: {args.provider if args.provider else 'auto-select'}")
    
    try:
        generator = BookGenerator(excel_path, provider=args.provider)
        
        # Override max_workers if specified
        if args.workers:
            generator.max_workers = args.workers
            generator.logger.info(f"Worker threads overridden to: {generator.max_workers}")
            
        generator.generate_book(output_dir=args.output)
    except Exception as e:
        logger.error(f"Error: {str(e)}")
        sys.exit(1)