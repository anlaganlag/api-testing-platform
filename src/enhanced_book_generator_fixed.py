import os
import sys
import json
import time
import random
import logging
import argparse
import concurrent.futures
from threading import Lock
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Tuple, Optional, Any, Callable
import hashlib

import pandas as pd
import requests
from tqdm import tqdm
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class APIClient:
    """A wrapper class for API clients with retry logic and fallback."""
    def __init__(self, provider_config: Dict[str, Any], logger: logging.Logger):
        self.provider_name = provider_config['name']
        self.api_key = provider_config['api_key']
        self.api_url = provider_config['api_url']
        self.model = provider_config['model']
        self.get_headers = provider_config['headers']
        self.format_request = provider_config['request_format']
        self.parse_response = provider_config['parse_response']
        self.enabled = provider_config['enabled']
        self.last_used = provider_config['last_used']
        self.success_count = provider_config['success_count']
        self.failure_count = provider_config['failure_count']
        self.consecutive_failures = provider_config['consecutive_failures']
        self.logger = logger
        self.lock = Lock()  # Lock for thread-safe updates to stats
        self.last_request_time = 0  # 记录上次请求时间，用于限制请求频率
        self.retry_delay = 2  # 初始重试延迟（秒）
        self.max_retries = 3  # 最大重试次数
    
    def make_api_call(self, prompt: str, max_tokens: int) -> Optional[str]:
        """Make an API call with error handling and retry logic."""
        # Check if this provider is enabled
        if not self.enabled:
            self.logger.warning(f"Provider {self.provider_name} is disabled due to previous failures")
            return None
        
        # Implement rate limiting
        self._respect_rate_limit()
        
        headers = self.get_headers(self.api_key)
        data = self.format_request(prompt, max_tokens)
        
        # For Gemini, we need to append the model and generateContent to the URL
        url = self.api_url
        params = {}
        if self.provider_name == 'gemini':
            url = f"{self.api_url}/{self.model}:generateContent"
            params = {"key": self.api_key}
            # Remove Authorization header for Gemini
            if "Authorization" in headers:
                del headers["Authorization"]
        
        self.logger.info(f"Making API call to provider: {self.provider_name}")
        self.logger.debug(f"Request URL: {url}")
        
        # Implement retry logic
        for retry_attempt in range(self.max_retries + 1):
            try:
                # Record request time
                self.last_request_time = time.time()
                start_time = time.time()
                response = requests.post(url, headers=headers, json=data, params=params, timeout=60)
                elapsed_time = time.time() - start_time
                
                # Handle different types of errors
                if response.status_code != 200:
                    error_msg = f"API error ({self.provider_name}, status {response.status_code}): {response.text[:200]}"
                    self.logger.error(error_msg)
                    
                    # Special error handling
                    if response.status_code == 401:
                        self.logger.error(f"Authentication error for {self.provider_name}. Check API key format.")
                        # Don't retry authentication errors
                        break
                    elif response.status_code == 429 or response.status_code == 422:
                        # Rate limit errors, increase retry delay
                        with self.lock:
                            self.failure_count += 1
                            self.consecutive_failures += 1
                            
                        # Disable provider if it has too many consecutive failures
                        if self.consecutive_failures >= 3:
                            self.logger.warning(f"Disabling provider {self.provider_name} due to {self.consecutive_failures} consecutive failures")
                            self.enabled = False
                            return None
                            
                        retry_delay = self._calculate_backoff_time(retry_attempt)
                        self.logger.warning(f"Rate limit exceeded for {self.provider_name}. Retrying in {retry_delay} seconds...")
                        time.sleep(retry_delay)
                        continue
                    elif response.status_code >= 500:
                        # Server errors, we can retry
                        with self.lock:
                            self.failure_count += 1
                            self.consecutive_failures += 1
                        
                        retry_delay = self._calculate_backoff_time(retry_attempt)
                        self.logger.warning(f"Server error for {self.provider_name}. Retrying in {retry_delay} seconds...")
                        time.sleep(retry_delay)
                        continue
                    
                    # Other errors, update stats but don't retry
                    with self.lock:
                        self.failure_count += 1
                        self.consecutive_failures += 1
                    
                    # If we still have retries left, continue trying
                    if retry_attempt < self.max_retries:
                        retry_delay = self._calculate_backoff_time(retry_attempt)
                        self.logger.warning(f"Retrying {self.provider_name} in {retry_delay} seconds (attempt {retry_attempt+1}/{self.max_retries})")
                        time.sleep(retry_delay)
                        continue
                    else:
                        # Disable provider if it has too many consecutive failures
                        if self.consecutive_failures >= 3:
                            self.logger.warning(f"Disabling provider {self.provider_name} due to {self.consecutive_failures} consecutive failures")
                            self.enabled = False
                        return None
                
                # Successfully received response
                result = response.json()
                content = self.parse_response(result)
                
                if not content:
                    self.logger.warning(f"Empty content returned from {self.provider_name}")
                    
                    # Update statistics
                    with self.lock:
                        self.failure_count += 1
                        self.consecutive_failures += 1
                    
                    # If we still have retries left, continue trying
                    if retry_attempt < self.max_retries:
                        retry_delay = self._calculate_backoff_time(retry_attempt)
                        self.logger.warning(f"Retrying {self.provider_name} in {retry_delay} seconds (attempt {retry_attempt+1}/{self.max_retries})")
                        time.sleep(retry_delay)
                        continue
                    else:
                        # Disable provider if it has too many consecutive failures
                        if self.consecutive_failures >= 3:
                            self.logger.warning(f"Disabling provider {self.provider_name} due to {self.consecutive_failures} consecutive failures")
                            self.enabled = False
                        return None
                
                # Update success statistics
                with self.lock:
                    self.last_used = time.time()
                    self.success_count += 1
                    self.consecutive_failures = 0
                
                self.logger.info(f"Successful API call to {self.provider_name} (took {elapsed_time:.2f}s)")
                return content
                    
            except Exception as e:
                self.logger.error(f"API call exception ({self.provider_name}): {str(e)}")
                
                # Update statistics
                with self.lock:
                    self.failure_count += 1
                    self.consecutive_failures += 1
                
                # If we still have retries left, continue trying
                if retry_attempt < self.max_retries:
                    retry_delay = self._calculate_backoff_time(retry_attempt)
                    self.logger.warning(f"Retrying {self.provider_name} in {retry_delay} seconds (attempt {retry_attempt+1}/{self.max_retries})")
                    time.sleep(retry_delay)
                    continue
                else:
                    # Disable provider if it has too many consecutive failures
                    if self.consecutive_failures >= 3:
                        self.logger.warning(f"Disabling provider {self.provider_name} due to {self.consecutive_failures} consecutive failures")
                        self.enabled = False
                    return None
        
        return None
    
    def _respect_rate_limit(self):
        """Ensure requests don't exceed API rate limits"""
        # Set different minimum request intervals for different providers
        min_request_interval = 1.0  # Default minimum interval is 1 second
        
        # Set specific intervals for different providers
        if self.provider_name == 'deepseek':
            min_request_interval = 2.0  # DeepSeek needs longer intervals
        elif self.provider_name == 'gemini':
            min_request_interval = 1.5
        
        # If time since last request is less than minimum interval, wait
        elapsed = time.time() - self.last_request_time
        if elapsed < min_request_interval:
            sleep_time = min_request_interval - elapsed
            self.logger.debug(f"Rate limiting: Waiting {sleep_time:.2f}s before next {self.provider_name} request")
            time.sleep(sleep_time)
    
    def _calculate_backoff_time(self, retry_attempt: int) -> float:
        """Calculate exponential backoff time"""
        # Base delay is 2 seconds, doubled for each retry, with some randomness
        backoff_time = self.retry_delay * (2 ** retry_attempt) + random.uniform(0, 1)
        return min(backoff_time, 60)  # Maximum delay of 60 seconds
    
    def get_stats(self) -> Dict[str, Any]:
        """Get current provider statistics."""
        with self.lock:
            return {
                'name': self.provider_name,
                'success_count': self.success_count,
                'failure_count': self.failure_count,
                'consecutive_failures': self.consecutive_failures,
                'last_used': self.last_used,
                'success_rate': self.success_count / max(1, self.success_count + self.failure_count)
            }

class ContentCache:
    """A disk-based cache for generated content to avoid unnecessary API calls."""
    def __init__(self, cache_dir: str = "temp/cache"):
        self.cache_dir = Path(cache_dir)
        self.cache_dir.mkdir(exist_ok=True, parents=True)
        self.lock = Lock()  # Lock for thread-safe file operations
    
    def _get_cache_key(self, prompt: str) -> str:
        """Generate a unique cache key for a prompt."""
        return hashlib.md5(prompt.encode('utf-8')).hexdigest()
    
    def get(self, prompt: str) -> Optional[str]:
        """Get content from cache if it exists."""
        cache_key = self._get_cache_key(prompt)
        cache_file = self.cache_dir / f"{cache_key}.txt"
        
        if cache_file.exists():
            with self.lock:
                try:
                    with open(cache_file, 'r', encoding='utf-8') as f:
                        return f.read()
                except Exception:
                    return None
        return None
    
    def set(self, prompt: str, content: str) -> None:
        """Save content to cache."""
        cache_key = self._get_cache_key(prompt)
        cache_file = self.cache_dir / f"{cache_key}.txt"
        
        with self.lock:
            try:
                with open(cache_file, 'w', encoding='utf-8') as f:
                    f.write(content)
            except Exception as e:
                print(f"Error saving to cache: {str(e)}")

class EnhancedBookGenerator:
    def __init__(self, excel_path, provider=None, max_workers=None):
        """Initialize the EnhancedBookGenerator with the path to the Excel outline.
        
        Args:
            excel_path: Path to the Excel outline file
            provider: API provider to use ('deepseek', 'gemini', 'openrouter', 'siliconflow', 'ark', 'all', or None for random)
            max_workers: Maximum number of worker threads to use (default: auto-configure based on providers)
        """
        # Load environment variables
        load_dotenv(override=True)  # Force reload and override existing env vars
        
        # Set up logging
        self._setup_logging()
        
        # Initialize content cache
        self.cache = ContentCache()
        
        # Available API providers configuration
        provider_configs = {
            'deepseek': {
                'name': 'deepseek',
                'api_key': os.getenv('DEEPSEEK_API_KEY'),
                'api_url': os.getenv('DEEPSEEK_API_URL'),
                'model': os.getenv('DEEPSEEK_API_MODEL'),
                'headers': lambda key: {"Authorization": f"Bearer {key}"},
                'request_format': self._format_openai_request,
                'parse_response': self._parse_openai_response,
                'enabled': bool(os.getenv('DEEPSEEK_API_KEY') and os.getenv('DEEPSEEK_API_URL') and os.getenv('DEEPSEEK_API_MODEL')),
                'last_used': 0,  # Timestamp of last successful use
                'success_count': 0,
                'failure_count': 0,
                'consecutive_failures': 0
            },
            'gemini': {
                'name': 'gemini',
                'api_key': os.getenv('GEMINI_API_KEY'),
                'api_url': os.getenv('GEMINI_API_URL'),
                'model': os.getenv('GEMINI_API_MODEL'),
                'headers': lambda key: {"Content-Type": "application/json"},
                'request_format': self._format_gemini_request,
                'parse_response': self._parse_gemini_response,
                'enabled': bool(os.getenv('GEMINI_API_KEY') and os.getenv('GEMINI_API_URL') and os.getenv('GEMINI_API_MODEL')),
                'last_used': 0,
                'success_count': 0,
                'failure_count': 0,
                'consecutive_failures': 0
            },
            'openrouter': {
                'name': 'openrouter',
                'api_key': os.getenv('OPENROUTER_API_KEY'),
                'api_url': os.getenv('OPENROUTER_API_URL'),
                'model': os.getenv('OPENROUTER_MODEL'),
                'headers': lambda key: {
                    "Authorization": f"Bearer {key}",
                    "HTTP-Referer": "https://github.com/user/repo"
                },
                'request_format': self._format_openai_request,
                'parse_response': self._parse_openai_response,
                'enabled': bool(os.getenv('OPENROUTER_API_KEY') and os.getenv('OPENROUTER_API_URL') and os.getenv('OPENROUTER_MODEL')),
                'last_used': 0,
                'success_count': 0,
                'failure_count': 0,
                'consecutive_failures': 0
            },
            'siliconflow': {
                'name': 'siliconflow',
                'api_key': os.getenv('SILICONFLOW_API_KEY'),
                'api_url': os.getenv('SILICONFLOW_API_URL'),
                'model': os.getenv('SILICONFLOW_MODEL'),
                'headers': lambda key: {"Authorization": f"Bearer {key}"},
                'request_format': self._format_openai_request,
                'parse_response': self._parse_openai_response,
                'enabled': bool(os.getenv('SILICONFLOW_API_KEY') and os.getenv('SILICONFLOW_API_URL') and os.getenv('SILICONFLOW_MODEL')),
                'last_used': 0,
                'success_count': 0,
                'failure_count': 0,
                'consecutive_failures': 0
            },
            'ark': {
                'name': 'ark',
                'api_key': os.getenv('ARK_API_KEY'),
                'api_url': os.getenv('ARK_API_URL'),
                'model': os.getenv('ARK_MODEL'),
                'headers': lambda key: {"Authorization": f"Bearer {key}"},
                'request_format': self._format_openai_request,
                'parse_response': self._parse_openai_response,
                'enabled': bool(os.getenv('ARK_API_KEY') and os.getenv('ARK_API_URL') and os.getenv('ARK_MODEL')),
                'last_used': 0,
                'success_count': 0,
                'failure_count': 0,
                'consecutive_failures': 0
            },
            'dashscope': {
                'name': 'dashscope',
                'api_key': os.getenv('DASHSCOPE_API_KEY'),
                'api_url': os.getenv('DASHSCOP_API_URL'),
                'model': os.getenv('DASHSCOP_MODEL'),
                'headers': lambda key: {"Authorization": f"Bearer {key}"},
                'request_format': self._format_openai_request,
                'parse_response': self._parse_openai_response,
                'enabled': bool(os.getenv('DASHSCOPE_API_KEY') and os.getenv('DASHSCOP_API_URL') and os.getenv('DASHSCOP_MODEL')),
                'last_used': 0,
                'success_count': 0,
                'failure_count': 0,
                'consecutive_failures': 0
            }
        }
        
        # Initialize API clients
        self.api_clients = {}
        self.available_providers = []
        for provider_name, config in provider_configs.items():
            if config['enabled']:
                self.api_clients[provider_name] = APIClient(config, self.logger)
                self.available_providers.append(provider_name)
        
        if not self.available_providers:
            self.logger.error("No API providers available. Please check your .env file.")
            raise ValueError("No API providers available. Please check your .env file.")
        
        self.logger.info(f"Available API providers: {', '.join(self.available_providers)}")
        
        # Select provider mode
        self.use_all_providers = (provider == 'all')
        if self.use_all_providers:
            self.logger.info("Using all available API providers with fallback")
        else:
            # Select specific provider or random
            if provider and provider in self.available_providers:
                self.active_provider = provider
            else:
                self.active_provider = random.choice(self.available_providers)
            self.logger.info(f"Selected primary API provider: {self.active_provider}")
        
        # Initialize other properties
        self.excel_path = excel_path
        self.outline = None
        self.generated_chapters = {}
        self.context = []
        self.doc = Document()
        self._setup_document_style()
        
        # 减少并发请求数量，避免触发API限制
        if max_workers is not None:
            self.max_workers = max_workers
        elif self.use_all_providers:
            # 减少并发数量，避免触发API限制
            self.max_workers = min(len(self.available_providers), 3)
        else:
            self.max_workers = 2  # 默认降低并发数
            
        self.logger.info(f"Using {self.max_workers} worker threads")
        
        # Optional configuration
        self.temperature = float(os.getenv('TEMPERATURE', '0.7'))
        self.max_tokens = int(os.getenv('MAX_TOKENS', '10000'))
        
    def _setup_logging(self):
        """Set up logging configuration."""
        # Create logs directory if it doesn't exist
        log_dir = Path("logs")
        log_dir.mkdir(exist_ok=True)
        
        # Configure logger
        self.logger = logging.getLogger('EnhancedBookGenerator')
        self.logger.setLevel(logging.INFO)
        
        # Clear existing handlers if any
        if self.logger.handlers:
            self.logger.handlers.clear()
        
        # Create file handler for logging to file
        log_file = log_dir / f"book_generator_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
        file_handler = logging.FileHandler(log_file, encoding='utf-8')
        file_handler.setLevel(logging.INFO)
        
        # Create console handler for logging to console
        console_handler = logging.StreamHandler()
        console_handler.setLevel(logging.INFO)
        
        # Create formatter and add it to the handlers
        formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
        file_handler.setFormatter(formatter)
        console_handler.setFormatter(formatter)
        
        # Add the handlers to the logger
        self.logger.addHandler(file_handler)
        self.logger.addHandler(console_handler)
        
        self.logger.info("Logging initialized")

    def _format_openai_request(self, prompt, max_tokens):
        """Format request data for OpenAI-compatible APIs (DeepSeek, OpenRouter, Siliconflow, Ark)."""
        return {
            "model": self.model if hasattr(self, 'model') else None,
            "messages": [
                {"role": "system", "content": "你是一位专业的教材作者，擅长创作清晰、专业的教育内容。"},
                {"role": "user", "content": prompt}
            ],
            "temperature": self.temperature,
            "max_tokens": max_tokens,
            "stream": False
        }
    
    def _format_gemini_request(self, prompt, max_tokens):
        """Format request data for Gemini API."""
        return {
            "contents": [
                {"role": "user", "parts": [{"text": prompt}]}
            ],
            "generationConfig": {
                "temperature": self.temperature,
                "maxOutputTokens": max_tokens
            }
        }
    
    def _parse_gemini_response(self, response_json):
        """Parse response from Gemini API."""
        try:
            return response_json.get('candidates', [{}])[0].get('content', {}).get('parts', [{}])[0].get('text', '')
        except (KeyError, IndexError, TypeError):
            self.logger.error(f"无法解析Gemini响应: {json.dumps(response_json, ensure_ascii=False)[:200]}...")
            return None
    
    def _parse_openai_response(self, response_json):
        """Parse response from OpenAI-compatible APIs."""
        try:
            return response_json.get('choices', [{}])[0].get('message', {}).get('content', '')
        except (KeyError, IndexError, TypeError):
            self.logger.error(f"无法解析OpenAI兼容响应: {json.dumps(response_json, ensure_ascii=False)[:200]}...")
            return None

    def _setup_document_style(self):
        """Set up the Word document style."""
        # Set up styles for different heading levels
        style = self.doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
        
        # Title style
        style = self.doc.styles['Title']
        style.font.name = '黑体'
        style.font.size = Pt(24)
        style.font.bold = True
        
        # Heading 1 style
        style = self.doc.styles['Heading 1']
        style.font.name = '黑体'
        style.font.size = Pt(18)
        style.font.bold = True
        
        # Heading 2 style
        style = self.doc.styles['Heading 2']
        style.font.name = '黑体'
        style.font.size = Pt(16)
        style.font.bold = True
        
        # 添加TOC样式
        if 'TOC 1' not in self.doc.styles:
            self.doc.styles.add_style('TOC 1', 1)
        if 'TOC 2' not in self.doc.styles:
            self.doc.styles.add_style('TOC 2', 1)

        # Add title page
        title = self.doc.add_paragraph("财富管理教材", style='Title')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph()  # Add some space
        self.doc.add_page_break()

    def load_outline(self):
        """Load and validate the book outline from Excel file."""
        try:
            self.logger.info(f"Loading outline from Excel file: {self.excel_path}")
            self.outline = pd.read_excel(self.excel_path)
            # Check if the file has the expected Chinese column names
            required_columns = ['课程模块', '课程主题', '课程大纲', '内容说明']
            if not all(col in self.outline.columns for col in required_columns):
                self.logger.error(f"Excel file must contain columns: {required_columns}")
                raise ValueError(f"Excel file must contain columns: {required_columns}")
            # Clean any NaN values
            self.outline = self.outline.fillna('')
            self.logger.info(f"Successfully loaded outline with {len(self.outline)} rows")
            return True
        except Exception as e:
            self.logger.error(f"Error loading outline: {str(e)}")
            return False

    def _select_best_provider(self, exclude_provider=None) -> str:
        """Select the best provider based on performance metrics."""
        # Only include enabled providers that aren't explicitly excluded
        available_providers = [
            p for p in self.available_providers 
            if (p != exclude_provider and self.api_clients[p].enabled)
        ]
        
        if not available_providers:
            return None
        
        # Sort providers by success rate and consecutive failures
        sorted_providers = sorted(
            available_providers,
            key=lambda p: (
                self.api_clients[p].consecutive_failures,  # Fewer consecutive failures preferred
                -self.api_clients[p].success_count / max(1, self.api_clients[p].success_count + self.api_clients[p].failure_count),  # Higher success rate preferred
                -(self.api_clients[p].last_used or 0)  # Recent unused providers preferred
            )
        )
        
        # Log provider selection criteria
        self.logger.info("Provider selection report:")
        for provider in sorted_providers:
            client = self.api_clients[provider]
            success_rate = client.success_count / max(1, client.success_count + client.failure_count)
            self.logger.info(f"  {provider}: 连续失败次数={client.consecutive_failures}, "
                           f"成功率={success_rate:.2f}, "
                           f"最近使用时间={client.last_used}")
        
        return sorted_providers[0] if sorted_providers else None

    def _make_api_call(self, prompt: str, max_tokens: int = 8000) -> Optional[str]:
        """Make an API call with caching, retry logic, and provider fallback."""
        # Check cache first
        cached_content = self.cache.get(prompt)
        if cached_content:
            self.logger.info("Using cached content")
            return cached_content
        
        # 设置重试参数
        max_retries = 3
        retry_count = 0
        
        # 如果使用所有提供商，尝试并行调用
        if self.use_all_providers:
            content = self._make_parallel_api_calls(prompt, max_tokens)
        else:
            # 使用选定的提供商并支持回退
            content = self._make_single_provider_call(self.active_provider, prompt, max_tokens, retry_count=retry_count, max_retries=max_retries)
        
        # 如果成功，缓存结果
        if content:
            self.cache.set(prompt, content)
        
        return content
    
    def _make_single_provider_call(self, provider_name: str, prompt: str, max_tokens: int, is_fallback: bool = False, retry_count: int = 0, max_retries: int = 3) -> Optional[str]:
        """Make an API call to a specific provider with fallback to others if it fails."""
        if provider_name not in self.api_clients:
            self.logger.error(f"Provider {provider_name} not available")
            return None
        
        client = self.api_clients[provider_name]
        content = client.make_api_call(prompt, max_tokens)
        
        if content:
            return content
        
        # 如果还有重试次数，先尝试重试当前提供商
        if retry_count < max_retries:
            # 计算指数退避时间：基础时间为2秒，每次重试翻倍，并添加随机性
            backoff_time = 2 * (2 ** retry_count) + random.uniform(0, 1)
            # 限制最大等待时间为60秒
            backoff_time = min(backoff_time, 60)
            
            self.logger.warning(f"API调用失败，将在 {backoff_time:.2f} 秒后重试 {provider_name} (尝试 {retry_count+1}/{max_retries})")
            time.sleep(backoff_time)
            
            # 递归调用自身进行重试，增加重试计数
            return self._make_single_provider_call(provider_name, prompt, max_tokens, is_fallback, retry_count + 1, max_retries)
        
        # 如果已经是回退调用，不再尝试其他提供商
        if is_fallback:
            return None
        
        # 尝试回退到另一个提供商
        fallback_provider = self._select_best_provider(exclude_provider=provider_name)
        if not fallback_provider:
            self.logger.error("No fallback providers available")
            return None
        
        self.logger.info(f"Falling back to provider: {fallback_provider}")
        return self._make_single_provider_call(fallback_provider, prompt, max_tokens, is_fallback=True)
    
    def _make_parallel_api_calls(self, prompt: str, max_tokens: int, retry_count: int = 0, max_retries: int = 3) -> Optional[str]:
        """Make parallel API calls to all available providers and return the first successful result."""
        self.logger.info(f"Making parallel API calls to {len(self.available_providers)} providers")
        
        # Sort providers by performance metrics
        sorted_providers = self._get_sorted_providers()
        
        # 限制并发请求数量，避免同时发送过多请求
        max_concurrent = min(len(sorted_providers), 2)  # 最多同时请求2个提供商
        
        # Use threading to make parallel API calls
        with concurrent.futures.ThreadPoolExecutor(max_workers=max_concurrent) as executor:
            # Submit API calls with staggered starts to avoid rate limits
            future_to_provider = {}
            for i, provider in enumerate(sorted_providers[:max_concurrent]):
                # 错开请求时间，避免同时发送
                if i > 0:
                    time.sleep(1.5)  # 每个请求间隔1.5秒
                future = executor.submit(self._make_single_provider_call, provider, prompt, max_tokens, True)
                future_to_provider[future] = provider
            
            # Get the first successful result
            for future in concurrent.futures.as_completed(future_to_provider):
                provider = future_to_provider[future]
                try:
                    result = future.result()
                    if result:
                        self.logger.info(f"Using result from {provider}")
                        return result
                except Exception as e:
                    self.logger.error(f"Error in parallel API call to {provider}: {str(e)}")
        
        self.logger.error("All parallel API calls failed")
        
        # 如果还有重试次数，使用指数退避算法进行重试
        if retry_count < max_retries:
            # 计算指数退避时间：基础时间为2秒，每次重试翻倍，并添加随机性
            backoff_time = 2 * (2 ** retry_count) + random.uniform(0, 1)
            # 限制最大等待时间为60秒
            backoff_time = min(backoff_time, 60)
            
            self.logger.warning(f"所有并行API调用失败，将在 {backoff_time:.2f} 秒后重试 (尝试 {retry_count+1}/{max_retries})")
            time.sleep(backoff_time)
            
            # 递归调用自身进行重试，增加重试计数
            return self._make_parallel_api_calls(prompt, max_tokens, retry_count + 1, max_retries)
        
        return None
    
    def _get_sorted_providers(self) -> List[str]:
        """Sort providers by performance metrics."""
        return sorted(
            self.available_providers,
            key=lambda p: (
                self.api_clients[p].consecutive_failures,  # 连续失败次数少的优先
                -self.api_clients[p].success_count / max(1, self.api_clients[p].success_count + self.api_clients[p].failure_count),  # 成功率高的优先
                -(self.api_clients[p].last_used or 0)  # 最近未使用的优先
            )
        )

    def generate_section_content(self, module: str, topic: str, outline_point: str, description: str) -> Optional[str]:
        """Generate content for a single section using the configured API providers."""
        # Create temp directory if it doesn't exist
        temp_dir = Path("temp/sections")
        temp_dir.mkdir(exist_ok=True, parents=True)
        
        # Generate a safe filename from the topic and outline point
        safe_filename = "".join(x for x in f"{topic}_{outline_point[:30]}" if x.isalnum() or x in (' ', '-', '_')).rstrip()
        temp_file = temp_dir / f"{safe_filename}.md"
        
        # Check if we already have this section generated
        if temp_file.exists():
            self.logger.info(f"Loading section from cache: {outline_point[:30]}...")
            with open(temp_file, 'r', encoding='utf-8') as f:
                content = f.read()
                if content.strip():  # If file is not empty
                    return content

        prompt = f"""请基于以下大纲点写一个详细的小节内容（约800-1000字）：

模块：{module}
主题：{topic}
大纲点：{outline_point}
详细说明：{description}

要求：
1. 内容应该具体且专业
2. 使用清晰的写作风格
3. 包含相关的例子和解释
4. 确保内容完整且自成一体
5. 目标字数800-1000字

请撰写小节内容："""

        try:
            self.logger.info(f"Generating section: {outline_point[:30]}...")
            section_content = self._make_api_call(prompt, max_tokens=2000)
            if not section_content:
                self.logger.error(f"Failed to generate content for section: {outline_point[:30]}...")
                return None

            # Save the generated content to temp file immediately with markdown formatting
            with open(temp_file, 'w', encoding='utf-8') as f:
                f.write(f"### {outline_point}\n\n")  # Add section heading
                f.write(section_content)
            self.logger.info(f"Section saved to temp file: {temp_file}")
            
            return section_content
        except Exception as e:
            self.logger.error(f"Error generating section content ({outline_point[:30]}...): {str(e)}")
            return None

    def stitch_chapter_content(self, module: str, topic: str, sections_content: List[str]) -> Optional[str]:
        """Stitch together section contents into a coherent chapter."""
        temp_dir = Path("temp/chapters")
        temp_dir.mkdir(exist_ok=True, parents=True)
        
        # Generate a safe filename from the topic
        safe_filename = "".join(x for x in topic if x.isalnum() or x in (' ', '-', '_')).rstrip()
        temp_file = temp_dir / f"{safe_filename}.md"
        
        # Check if we already have this chapter stitched
        if temp_file.exists():
            self.logger.info(f"Loading chapter from cache: {topic}")
            with open(temp_file, 'r', encoding='utf-8') as f:
                content = f.read()
                if content.strip():  # If file is not empty
                    return content
        
        # If no sections content provided, return None
        if not sections_content:
            self.logger.error(f"No sections content provided for chapter: {topic}")
            return None
        
        # Combine all sections into a single chapter
        chapter_content = f"## {topic}\n\n"
        for section in sections_content:
            if section:
                chapter_content += section + "\n\n"
        
        # Save the stitched chapter to temp file
        with open(temp_file, 'w', encoding='utf-8') as f:
            f.write(chapter_content)
        self.logger.info(f"Chapter saved to temp file: {temp_file}")
        
        return chapter_content

    def generate_chapter(self, module: str, topic: str, outline_points: List[Tuple[str, str]]) -> Optional[str]:
        """Generate a complete chapter with multiple sections."""
        self.logger.info(f"Starting chapter generation: {topic} with {len(outline_points)} outline points")
        
        # Generate each section in parallel
        self.logger.info(f"Generating {len(outline_points)} sections in parallel for topic: {topic}")
        sections_content = []
        
        # 减少并发生成的数量，改为顺序生成以避免API限制
        if len(outline_points) > 3:
            self.logger.info(f"Large chapter detected ({len(outline_points)} sections). Generating in batches.")
            batches = [outline_points[i:i+2] for i in range(0, len(outline_points), 2)]
            
            for batch_idx, batch in enumerate(batches):
                self.logger.info(f"Processing batch {batch_idx+1}/{len(batches)} with {len(batch)} sections")
                batch_results = []
                
                with concurrent.futures.ThreadPoolExecutor(max_workers=min(len(batch), 2)) as executor:
                    future_to_point = {}
                    for i, (outline_point, description) in enumerate(batch):
                        # 错开请求时间
                        if i > 0:
                            time.sleep(2)
                        future = executor.submit(self.generate_section_content, module, topic, outline_point, description)
                        future_to_point[future] = outline_point
                    
                    for future in concurrent.futures.as_completed(future_to_point):
                        outline_point = future_to_point[future]
                        try:
                            section_content = future.result()
                            if section_content:
                                batch_results.append(section_content)
                            else:
                                self.logger.warning(f"Failed to generate section: {outline_point[:30]}...")
                        except Exception as e:
                            self.logger.error(f"Error in section generation: {str(e)}")
                
                sections_content.extend(batch_results)
                # 批次之间添加延迟，避免触发API限制
                if batch_idx < len(batches) - 1:
                    delay = 5  # 5秒延迟
                    self.logger.info(f"Waiting {delay} seconds before next batch...")
                    time.sleep(delay)
        else:
            # 对于小章节，可以并行生成
            with concurrent.futures.ThreadPoolExecutor(max_workers=min(len(outline_points), 2)) as executor:
                future_to_point = {}
                for i, (outline_point, description) in enumerate(outline_points):
                    # 错开请求时间
                    if i > 0:
                        time.sleep(2)
                    future = executor.submit(self.generate_section_content, module, topic, outline_point, description)
                    future_to_point[future] = outline_point
                
                for future in concurrent.futures.as_completed(future_to_point):
                    outline_point = future_to_point[future]
                    try:
                        section_content = future.result()
                        if section_content:
                            sections_content.append(section_content)
                        else:
                            self.logger.warning(f"Failed to generate section: {outline_point[:30]}...")
                    except Exception as e:
                        self.logger.error(f"Error in section generation: {str(e)}")
        
        self.logger.info(f"Completed generating {len(sections_content)}/{len(outline_points)} sections for topic: {topic}")
        
        # If no sections were generated successfully, return None
        if not sections_content:
            self.logger.error(f"Error: All sections failed for chapter {topic}")
            return None
        
        # Stitch together the sections into a coherent chapter
        return self.stitch_chapter_content(module, topic, sections_content)

    def generate_book(self):
        """Generate the complete book based on the Excel outline."""
        if not self.load_outline():
            self.logger.error("Failed to load outline. Exiting.")
            return False
        
        # Group outline by module and topic
        grouped_outline = {}
        for _, row in self.outline.iterrows():
            module = row['课程模块']
            topic = row['课程主题']
            outline_point = row['课程大纲']
            description = row['内容说明']
            
            if module not in grouped_outline:
                grouped_outline[module] = {}
            
            if topic not in grouped_outline[module]:
                grouped_outline[module][topic] = []
            
            grouped_outline[module][topic].append((outline_point, description))
        
        # Generate each chapter
        total_chapters = sum(len(topics) for topics in grouped_outline.values())
        self.logger.info(f"Starting book generation with {total_chapters} chapters")
        
        chapter_count = 0
        for module, topics in grouped_outline.items():
            self.logger.info(f"Processing module: {module} with {len(topics)} topics")
            
            # Add module heading to the document
            self.doc.add_heading(module, level=1)
            
            for topic, outline_points in topics.items():
                chapter_count += 1
                self.logger.info(f"Generating chapter {chapter_count}/{total_chapters}: {topic}")
                
                try:
                    # Generate the chapter content
                    chapter_content = self.generate_chapter(module, topic, outline_points)
                    
                    if not chapter_content:
                        self.logger.error(f"Failed to generate chapter {chapter_count}/{total_chapters}: {topic}")
                        continue
                    
                    # Add the chapter to the document
                    self.doc.add_heading(topic, level=2)
                    
                    # Parse the markdown content and add to document
                    lines = chapter_content.split('\n')
                    for line in lines:
                        # Skip the chapter heading (already added above)
                        if line.startswith('## '):
                            continue
                        # Handle section headings
                        elif line.startswith('### '):
                            self.doc.add_heading(line[4:], level=3)
                        # Handle regular paragraphs
                        elif line.strip():
                            self.doc.add_paragraph(line)
                    
                    # Add a page break after each chapter
                    self.doc.add_page_break()
                    
                    self.logger.info(f"Successfully added chapter {chapter_count}/{total_chapters}: {topic}")
                    
                except Exception as e:
                    self.logger.error(f"Error generating chapter {chapter_count}/{total_chapters}: {topic} - {str(e)}")
        
        # Save the document
        output_dir = Path("output")
        output_dir.mkdir(exist_ok=True)
        output_file = output_dir / "完整教材.docx"
        self.doc.save(output_file)
        self.logger.info(f"Book generation complete. Saved to {output_file}")
        
        # Save metadata
        metadata = {
            "title": "财富管理教材",
            "chapters": chapter_count,
            "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "api_providers": self.available_providers
        }
        with open(output_dir / "book_metadata.json", 'w', encoding='utf-8') as f:
            json.dump(metadata, f, ensure_ascii=False, indent=2)
        
        return True

    def generate_sample_chapter(self, chapter_index=0):
        """Generate a sample chapter for testing."""
        if not self.load_outline():
            self.logger.error("Failed to load outline. Exiting.")
            return False
        
        # Group outline by module and topic
        grouped_outline = {}
        for _, row in self.outline.iterrows():
            module = row['课程模块']
            topic = row['课程主题']
            outline_point = row['课程大纲']
            description = row['内容说明']
            
            if module not in grouped_outline:
                grouped_outline[module] = {}
            
            if topic not in grouped_outline[module]:
                grouped_outline[module][topic] = []
            
            grouped_outline[module][topic].append((outline_point, description))
        
        # Flatten the grouped outline into a list of (module, topic, outline_points) tuples
        chapters = []
        for module, topics in grouped_outline.items():
            for topic, outline_points in topics.items():
                chapters.append((module, topic, outline_points))
        
        if not chapters:
            self.logger.error("No chapters found in outline.")
            return False
        
        # Select the chapter to generate
        if chapter_index >= len(chapters):
            self.logger.warning(f"Chapter index {chapter_index} out of range. Using first chapter.")
            chapter_index = 0
        
        module, topic, outline_points = chapters[chapter_index]
        self.logger.info(f"Generating sample chapter: {topic} with {len(outline_points)} outline points")
        
        # Generate the chapter content
        chapter_content = self.generate_chapter(module, topic, outline_points)
        
        if not chapter_content:
            self.logger.error(f"Failed to generate sample chapter: {topic}")
            return False
        
        # Save the chapter as a markdown file
        output_dir = Path("output")
        output_dir.mkdir(exist_ok=True)
        safe_filename = "".join(x for x in topic if x.isalnum() or x in (' ', '-', '_')).rstrip()
        output_file = output_dir / f"{safe_filename}.md"
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(chapter_content)
        
        self.logger.info(f"Sample chapter generation complete. Saved to {output_file}")
        return True

# Main execution
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate a book from an Excel outline using AI.")
    parser.add_argument("--excel", "-e", default="data/book_outline.xlsx", help="Path to Excel outline file")
    parser.add_argument("--provider", "-p", choices=["deepseek", "gemini", "openrouter", "siliconflow", "ark", "dashscope", "all"], 
                        default="all", help="API provider to use (default: all)")
    parser.add_argument("--sample", "-s", type=int, help="Generate a sample chapter (specify index)")
    parser.add_argument("--threads", "-t", type=int, help="Maximum number of worker threads")
    
    args = parser.parse_args()
    
    try:
        generator = EnhancedBookGenerator(args.excel, provider=args.provider, max_workers=args.threads)
        
        if args.sample is not None:
            generator.generate_sample_chapter(args.sample)
        else:
            generator.generate_book()
            
    except Exception as e:
        logging.error(f"Error: {str(e)}")
        sys.exit(1)